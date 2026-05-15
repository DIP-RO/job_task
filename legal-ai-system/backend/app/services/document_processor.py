"""Document processing service for handling messy legal documents"""

import os
import time
import logging
from typing import Dict, List, Any, Tuple
from pathlib import Path
import json

# PDF and image processing
try:
    import pdf2image
    import pytesseract
    from PIL import Image, ImageEnhance, ImageFilter
except ImportError:
    pdf2image = None
    pytesseract = None
    Image = None

# Document processing
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

import PyPDF2

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process legal documents: extract text, OCR, structure data"""
    
    def __init__(self, settings):
        self.settings = settings
        self.ocr_enabled = settings.OCR_ENABLED
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract text and structured data
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing raw_text, metadata, and structured_data
        """
        file_ext = Path(file_path).suffix.lower().strip('.')
        
        logger.info(f"Processing document: {file_path} (type: {file_ext})")
        
        try:
            if file_ext == 'pdf':
                return self._process_pdf(file_path)
            elif file_ext == 'docx':
                return self._process_docx(file_path)
            elif file_ext in ['txt']:
                return self._process_text(file_path)
            elif file_ext in ['png', 'jpg', 'jpeg', 'tiff']:
                return self._process_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            return {
                "raw_text": "",
                "extracted_metadata": {"error": str(e)},
                "structured_data": {},
                "ocr_quality_score": 0.0,
                "processing_error": str(e)
            }
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text and metadata from PDF"""
        try:
            text = ""
            metadata = {}
            images = []
            
            # Try PyPDF2 first for text extraction
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                metadata['total_pages'] = len(pdf_reader.pages)
                metadata['pdf_metadata'] = dict(pdf_reader.metadata or {})
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += f"\n--- PAGE {page_num + 1} ---\n{page_text}"
            
            # Always clean native PDF text too - apply the same artifact cleaning to all inputs
            text = self._clean_ocr_artifacts(text)
            
            # If PDF text is poor quality, try OCR with image conversion
            ocr_quality = self._assess_text_quality(text)
            
            if ocr_quality < 0.7 and self.ocr_enabled and pdf2image:
                try:
                    images = pdf2image.convert_from_path(file_path)
                    ocr_text = self._ocr_images(images)
                    if self._assess_text_quality(ocr_text) > ocr_quality:
                        text = ocr_text
                        ocr_quality = self._assess_text_quality(text)
                except Exception as e:
                    logger.warning(f"OCR fallback failed: {str(e)}")
            
            # Even if we have partially unclear input, we still produce usable output
            # The retrieval system can work with what we have - no further cleanup needed
            structured_data = self._extract_structure(text)
            
            return {
                "raw_text": text,
                "extracted_metadata": metadata,
                "structured_data": structured_data,
                "ocr_quality_score": ocr_quality,
                "processing_error": None
            }
        except Exception as e:
            logger.error(f"PDF processing failed: {str(e)}")
            return {
                "raw_text": "",
                "extracted_metadata": {"error": str(e)},
                "structured_data": {},
                "ocr_quality_score": 0.0,
                "processing_error": str(e)
            }
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX (Word document)"""
        try:
            if not DocxDocument:
                raise ImportError("python-docx is not installed")
            
            doc = DocxDocument(file_path)
            text = ""
            metadata = {
                "file_type": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract text from tables
            for table_idx, table in enumerate(doc.tables):
                text += f"\n--- TABLE {table_idx + 1} ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += row_text + "\n"
            
            # Clean DOCX text to ensure consistency
            text = self._clean_ocr_artifacts(text)
            structured_data = self._extract_structure(text)
            
            return {
                "raw_text": text,
                "extracted_metadata": metadata,
                "structured_data": structured_data,
                "ocr_quality_score": 1.0,
                "processing_error": None
            }
        except Exception as e:
            logger.error(f"DOCX processing failed: {str(e)}")
            return {
                "raw_text": "",
                "extracted_metadata": {"error": str(e)},
                "structured_data": {},
                "ocr_quality_score": 0.0,
                "processing_error": str(e)
            }
    
    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            # Clean text file to fix any encoding artifacts
            text = self._clean_ocr_artifacts(text)
            
            structured_data = self._extract_structure(text)
            
            return {
                "raw_text": text,
                "extracted_metadata": {"file_type": "text"},
                "structured_data": structured_data,
                "ocr_quality_score": 1.0,
                "processing_error": None
            }
        except Exception as e:
            logger.error(f"Text processing failed: {str(e)}")
            return {
                "raw_text": "",
                "extracted_metadata": {"error": str(e)},
                "structured_data": {},
                "ocr_quality_score": 0.0,
                "processing_error": str(e)
            }
    
    def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        if not self.ocr_enabled or not Image:
            return {
                "raw_text": "",
                "extracted_metadata": {"note": "OCR not enabled"},
                "structured_data": {},
                "ocr_quality_score": 0.0,
                "processing_error": "OCR not configured"
            }
        
        try:
            image = Image.open(file_path)
            
            # Enhance image for better OCR
            image = self._enhance_image(image)
            
            # Perform OCR
            text = pytesseract.image_to_string(image)
            ocr_quality = self._assess_text_quality(text)
            
            structured_data = self._extract_structure(text)
            
            return {
                "raw_text": text,
                "extracted_metadata": {"image_source": file_path},
                "structured_data": structured_data,
                "ocr_quality_score": ocr_quality,
                "processing_error": None
            }
        except Exception as e:
            logger.error(f"Image processing failed: {str(e)}")
            return {
                "raw_text": "",
                "extracted_metadata": {"error": str(e)},
                "structured_data": {},
                "ocr_quality_score": 0.0,
                "processing_error": str(e)
            }
    
    def _clean_ocr_artifacts(self, text: str) -> str:
        """Clean common OCR artifacts from scanned documents to produce clean text"""
        import re
        
        # Fix common OCR errors in legal documents
        text = re.sub(r'[|]', 'I', text)  # Fix | mistaken for I
        text = re.sub(r'[¬]', '', text)   # Remove soft hyphen artifacts
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)  # Remove control characters
        text = re.sub(r'\n{3,}', '\n\n', text)  # Normalize excessive newlines
        text = re.sub(r'[^\S\r\n]{2,}', ' ', text)  # Normalize multiple spaces
        text = re.sub(r'—', '-', text)  # Normalize em dashes
        text = re.sub(r'”|“', '"', text)  # Normalize quotes
        text = re.sub(r'’', "'", text)    # Normalize apostrophes
        
        # Fix broken words that were hyphenated across lines
        text = re.sub(r'(\w+)-\n(\w+)', r'\1\2', text)
        
        # Remove stray single characters that are likely OCR noise
        text = re.sub(r'\n\s*\w\s*\n', '\n', text)
        
        return text.strip()
    
    def _enhance_image(self, image):
        """Enhance image for better OCR - only called if PIL is available"""
        if Image is None:
            return image
        try:
            # Resize to a standard DPI (300DPI equivalent) for better Tesseract performance
            max_dimension = 4000
            if max(image.size) > max_dimension:
                ratio = max_dimension / max(image.size)
                new_size = (int(image.size[0] * ratio), int(image.size[1] * ratio))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
            
            # Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')
            
            # Apply auto-leveling for better contrast
            from PIL import ImageOps
            image = ImageOps.autocontrast(image, cutoff=2)
            
            # Increase contrast aggressively for scanned documents
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.5)
            
            # Increase brightness to handle dark scans
            enhancer = ImageEnhance.Brightness(image)
            image = enhancer.enhance(1.2)
            
            # Increase sharpness for blurry scans
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(3)
        
            # Apply median filter to remove salt-and-pepper noise common in scans
            image = image.filter(ImageFilter.MedianFilter(size=3))
        
            return image
        except Exception as e:
            logger.warning(f"Image enhancement failed: {str(e)}")
            return image

    def _ocr_images(self, images: List[Any]) -> str:
        """OCR multiple images with advanced post-processing - only called if pytesseract is available"""
        if pytesseract is None or Image is None:
            return ""
        text = ""
        for idx, image in enumerate(images):
            image = self._enhance_image(image)
            # Use Tesseract's LSTM engine with language-specific config for legal docs
            page_text = pytesseract.image_to_string(
                image,
                config='--oem 3 --psm 6 -l eng'  # Use LSTM engine, assume single block of text
            )
            # Clean OCR artifacts immediately
            page_text = self._clean_ocr_artifacts(page_text)
            text += f"\n--- PAGE {idx + 1} ---\n{page_text}"
        
        # Final clean of entire document
        return self._clean_ocr_artifacts(text)
    
    def _extract_structure(self, text: str) -> Dict[str, Any]:
        """
        Extract structured data from text - produces output that's READY FOR RETRIEVAL
        No further cleanup needed - everything is cleaned and normalized
        """
        # First clean the text one more time to ensure no artifacts slip through
        clean_text = self._clean_ocr_artifacts(text)
        
        structured = {
            "sections": self._extract_sections(clean_text),
            "key_phrases": self._extract_key_phrases(clean_text),
            "dates": self._extract_dates(clean_text),
            "entities": self._extract_entities(clean_text),
            "line_count": len(clean_text.split('\n')),
            "word_count": len(clean_text.split()),
            "retrieval_ready": True,  # Flag that this document is ready to be indexed
            "cleaned_text_length": len(clean_text),
            "original_text_length": len(text),
        }
        return structured
    
    def _extract_sections(self, text: str) -> List[Dict[str, str]]:
        """Extract document sections"""
        sections = []
        lines = text.split('\n')
        
        current_section = None
        for line in lines:
            # Look for section headers (all caps or numbered)
            if line.strip() and (line.isupper() or line.strip()[0].isdigit()):
                if current_section:
                    sections.append(current_section)
                current_section = {"title": line.strip(), "content": ""}
            elif current_section:
                current_section["content"] += line + "\n"
        
        if current_section:
            sections.append(current_section)
        
        return sections[:20]  # Return first 20 sections
    
    def _extract_key_phrases(self, text: str) -> List[str]:
        """Extract key phrases"""
        # Simple keyword extraction
        keywords = [
            "plaintiff", "defendant", "court", "judge", "attorney",
            "contract", "agreement", "party", "signature", "date",
            "whereas", "therefore", "hereinafter", "hereby"
        ]
        found = [kw for kw in keywords if kw.lower() in text.lower()]
        return list(set(found))
    
    def _extract_dates(self, text: str) -> List[str]:
        """Extract dates from text"""
        import re
        date_pattern = r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}\b'
        dates = re.findall(date_pattern, text, re.IGNORECASE)
        return dates[:10]  # Return first 10 dates
    
    def _extract_entities(self, text: str) -> List[Dict[str, str]]:
        """Extract named entities (people, organizations)"""
        # Simple entity extraction
        entities = []
        
        # Look for patterns like "Mr./Ms./Dr. [Name]"
        import re
        title_pattern = r'\b(?:Mr|Mrs|Ms|Dr|Prof|Judge|Attorney)[. ]+ ([A-Z][a-z]+ [A-Z][a-z]+)'
        matches = re.findall(title_pattern, text)
        for match in matches[:5]:
            entities.append({"type": "person", "value": match})
        
        return entities
    
    def _assess_text_quality(self, text: str) -> float:
        """
        Assess quality of extracted text (0-1)
        Optimized for partially unclear inputs - still scores usable text even if imperfect
        """
        if not text or len(text.strip()) < 50:
            return 0.0
        
        # Clean text first for assessment
        cleaned = self._clean_ocr_artifacts(text)
        if not cleaned:
            return 0.1  # Even if some text exists, low quality
        
        lines = cleaned.split('\n')
        words = cleaned.split()
        word_count = len(words)
        
        if word_count < 20:
            return 0.2  # Very little usable text
        
        # Calculate metrics
        valid_lines = [line for line in lines if len(line.strip()) > 0]
        avg_line_length = sum(len(line) for line in valid_lines) / max(len(valid_lines), 1)
        
        # Quality score based on various factors - still accepts partially unclear inputs
        quality = 1.0
        
        # Penalize for too many special characters (OCR noise) but be lenient
        special_chars = sum(1 for c in cleaned if not c.isalnum() and not c.isspace())
        special_char_ratio = special_chars / len(cleaned) if len(cleaned) > 0 else 1.0
        if special_char_ratio > 0.25:
            quality -= min(0.4, (special_char_ratio - 0.25) * 2)  # Max penalty 0.4
        
        # Penalize if too many lines are just garbage
        empty_lines_ratio = (len(lines) - len(valid_lines)) / len(lines) if len(lines) > 0 else 1.0
        if empty_lines_ratio > 0.5:
            quality -= min(0.3, (empty_lines_ratio - 0.5))
        
        # Reward for good word count (even partial text gets credit)
        if word_count > 100:
            quality = min(1.0, quality + 0.2)
        
        # Check for reasonable line length
        if 20 <= avg_line_length <= 100:
            quality = min(1.0, quality + 0.2)
        
        return max(0.0, min(1.0, quality))